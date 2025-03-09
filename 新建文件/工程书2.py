import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def replace_in_docx(doc_path, replacements):
    if not os.path.exists(doc_path) or not doc_path.lower().endswith('.docx'):
        print(f"文件不存在或不是有效的.docx文件: {doc_path}")
        return

    try:
        doc = Document(doc_path)
        modified = False

        # 遍历文档中的所有段落
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    # 替换文本并保留其他内容
                    new_paragraph_content = paragraph.text.replace(old_text, new_text)
                    paragraph.clear()  # 清空段落内容
                    paragraph.add_run(new_paragraph_content)  # 添加新内容
                    modified = True

                    # 设置字体和字号
                    for run in paragraph.runs:
                        run.font.name = '仿宋_GB2312'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(15)

        # 遍历文档中的所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                # 替换文本并保留其他内容
                                new_paragraph_content = paragraph.text.replace(old_text, new_text)
                                paragraph.clear()  # 清空段落内容
                                paragraph.add_run(new_paragraph_content)  # 添加新内容
                                modified = True

                                # 设置字体和字号
                                for run in paragraph.runs:
                                    run.font.name = '仿宋_GB2312'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                    run.font.size = Pt(15)

        if modified:
            doc.save(doc_path)
            print(f"文档 {doc_path} 已修改并保存。")
        else:
            print(f"文档 {doc_path} 未找到需要替换的内容，未修改。")
    except Exception as e:
        print(f"处理文档 {doc_path} 时发生错误：{e}")

def process_folder(folder_path, excel_path):
    workbook = load_workbook(excel_path)
    sheet = workbook['Sheet3']
    project_names = []
    values = []

    for row in sheet.iter_rows(min_row=2, values_only=True):
        project_names.append(row[3])  # 确保 D 列的索引是 3
        values.append(row[2])  # 确保 C 列的索引是 2

    print("项目名称列表:", project_names)
    print("数值列表:", values)

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.docx'):
                doc_path = os.path.join(root, file)

                if not project_names or not values:
                    print("Excel 数据不足，无法继续处理。")
                    break

                replacements = {
                    "湖北天门10kV天门市公安局业扩配套工程（项目编号1815J8240002）": project_names.pop(0),
                    "1237.00": values.pop(0)
                }
                replace_in_docx(doc_path, replacements)

root_folder = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\项目工程"
excel_path = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\模板.xlsx"
process_folder(root_folder, excel_path)