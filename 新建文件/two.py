import pandas as pd
from docx import Document
import os
import shutil

# === 配置参数 ===
src_word_path = r'C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\施工单位重要事项承诺书（紫电）.docx'  # 源Word文件路径
target_project_dir = r'C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\项目工程'  # 目标工程根目录
excel_data_path = r'C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\模板.xlsx'  # Excel数据文件路径

# === 读取Excel数据 ===
try:
    df = pd.read_excel(excel_data_path, sheet_name='Sheet3', header=None)
except Exception as e:
    print(f"读取Excel失败: {e}")
    exit()

# 提取C列(D列是Excel的第4列，索引为3)和D列数据
c_column = df.iloc[:, 2]  # C列对应索引2
d_column = df.iloc[:, 3]  # D列对应索引3

# 获取所有子文件夹列表
subfolders = []
for item in os.listdir(target_project_dir):
    if os.path.isdir(os.path.join(target_project_dir, item)):
        subfolders.append(item)

# 数据校验：子文件夹数量必须等于有效数据行数
valid_data_rows = len(c_column.dropna())  # 忽略C列为空的行
if len(subfolders) != valid_data_rows:
    print(f"数据异常：子文件夹数量({len(subfolders)})与有效数据行数({valid_data_rows})不一致")
    exit()

# === 处理每个子文件夹 ===
for idx, folder_name in enumerate(subfolders):
    current_dir = os.path.join(target_project_dir, folder_name)

    # 复制源Word文件到当前目录
    dest_word = os.path.join(current_dir, os.path.basename(src_word_path))
    if not os.path.exists(dest_word):
        try:
            shutil.copy2(src_word_path, dest_word)
            print(f"已复制文件到: {dest_word}")
        except Exception as e:
            print(f"复制文件失败: {e}")
            continue

    # 加载Word文档
    try:
        doc = Document(dest_word)
    except Exception as e:
        print(f"打开文档失败: {e}")
        continue

    # 获取替换数据
    replace_num = c_column[idx] if idx < len(c_column) else ""
    replace_project = d_column[idx] if idx < len(d_column) else ""

    # 执行全局替换（段落级别）
    for para in doc.paragraphs:
        para.text = para.text.replace('1237.00', replace_num)
        para.text = para.text.replace(
            '湖北天门10kV天门市公安局业扩配套工程（项目编号1815J8240002）',
            replace_project
        )

    # 如果需要处理表格中的内容，可以添加以下代码：
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             cell.text = cell.text.replace('1237.00', replace_num)
    #             cell.text = cell.text.replace(
    #                 '湖北天门10kV天门市公安局业扩配套工程（项目编号1815J8240002）',
    #                 replace_project
    #             )

    # 保存修改后的文档
    try:
        doc.save(dest_word)
        print(f"已成功处理: {dest_word}")
    except Exception as e:
        print(f"保存文档失败: {e}")

print("所有操作完成！")