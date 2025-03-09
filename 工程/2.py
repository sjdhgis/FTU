from docx import Document
import os
import shutil
from openpyxl import load_workbook


def safe_open_document(doc_path):
    """
    安全地打开一个Word文档，如果文件不存在或无法打开，则返回None。
    """
    try:
        return Document(doc_path)
    except Exception as e:
        print(f"无法打开文件：{doc_path}。错误：{e}")
        return None


def replace_text_in_docx(doc_path, replacements):
    """
    在指定的Word文档中替换文本内容。
    """
    doc = safe_open_document(doc_path)
    if doc is None:
        return

    modified = False
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                modified = True

    # 遍历文档中的所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
                            modified = True

    # 如果文档被修改，则保存
    if modified:
        doc.save(doc_path)
        print(f"文档 {doc_path} 已修改并保存。")
    else:
        print(f"文档 {doc_path} 未找到需要替换的内容，未修改。")
    doc.close()  # 确保文档被关闭


def batch_replace_in_docx_files(source_folder, target_folder, excel_path):
    """
    批量处理Word文档，根据Excel表格内容进行替换。
    """
    # 确保目标文件夹存在
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)

    # 加载Excel文件并读取Sheet3的C列和D列内容
    workbook = load_workbook(excel_path)
    sheet = workbook['Sheet3']
    project_names = []  # 项目名称列表
    values = []  # 数值列表

    for row in sheet.iter_rows(min_row=2, values_only=True):  # 假设第一行是标题行
        project_names.append(row[3])  # 项目名称在D列
        values.append(row[2])  # 数值在C列

    # 获取源文件夹中的所有.docx文件
    docx_files = [f for f in os.listdir(source_folder) if f.endswith('.docx')]

    # 检查文件数量是否匹配
    if len(docx_files) != len(project_names):
        print("警告：源文件夹中的Word文件数量与Excel的A列条目数量不匹配。")
        return

    # 复制并重命名文件
    for i, filename in enumerate(docx_files):
        src_path = os.path.join(source_folder, filename)
        new_name = project_names[i] + '.docx'  # 构造新文件名
        dst_path = os.path.join(target_folder, new_name)

        # 复制文件到目标文件夹并重命名
        shutil.copy(src_path, dst_path)
        print(f"文件已复制到 {dst_path}")

        # 替换文档中的内容
        replacements = {
            "湖北天门10kV天门市公安局业扩配套工程（项目编号1815J8240002）": project_names[i],
            "1237.00": values[i]
        }
        replace_text_in_docx(dst_path, replacements)


# 示例用法
source_folder = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\工程\1"  # 源文件夹路径
target_folder = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\工程\2"  # 目标文件夹路径
excel_path = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\模板.xlsx"  # Excel文件路径

batch_replace_in_docx_files(source_folder, target_folder, excel_path)