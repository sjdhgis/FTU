import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml

def update_word_file(file_path, replacements, font_name='仿宋_GB2312', font_size=Pt(10.5)):
    doc = Document(file_path)
    modified = False

    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
                modified = True
                for run in paragraph.runs:
                    run.font.name = font_name
                    run.font.size = font_size

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)
                            modified = True
                            for run in paragraph.runs:
                                run.font.name = font_name
                                run.font.size = font_size

    if modified:
        doc.save(file_path)
        print(f"文件 {file_path} 已更新。")
    else:
        print(f"文件 {file_path} 未找到需要替换的内容，无需更新。")

def batch_update_docx_files_from_excel(source_folder, excel_path):
    workbook = load_workbook(excel_path)
    sheet = workbook['Sheet3']  # 假设数据在Sheet3中

    project_names = sheet['D']  # 项目名称在D列
    values = sheet['C']  # 数值在C列

    for root, dirs, files in os.walk(source_folder):
        for file in files:
            if file.endswith('.docx'):
                file_path = os.path.join(root, file)
                replacements = {
                    "湖北天门10kV天门市公安局业扩配套工程（项目编号1815J8240002）": project_names.pop(0) if project_names else "",
                    "1237.00": values.pop(0) if values else ""
                }
                update_word_file(file_path, replacements)

# 示例用法
source_folder = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\工程\1"  # 源文件夹路径
excel_path = r"C:\Users\33\Desktop\2024年网改项目结算\网改项目\项目\模板.xlsx"  # Excel文件路径

batch_update_docx_files_from_excel(source_folder, excel_path)